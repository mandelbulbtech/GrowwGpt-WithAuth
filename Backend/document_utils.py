import os
import json
import tempfile
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_text_from_document(file_obj, file_extension):
    """Extract text from uploaded documents based on file type."""
    try:
        file_extension = file_extension.lower()
       
        # Basic text files - read directly
        if file_extension in ['.txt', '.md']:
            content = file_obj.read()
            # Try to decode as UTF-8, fallback to latin-1 if needed
            try:
                return content.decode('utf-8')
            except UnicodeDecodeError:
                return content.decode('latin-1')
       
        # JSON files - parse and format
        elif file_extension == '.json':
            content = file_obj.read()
            try:
                json_content = json.loads(content)
                return json.dumps(json_content, indent=2)
            except:
                # If JSON parsing fails, return raw content
                try:
                    return content.decode('utf-8')
                except UnicodeDecodeError:
                    return content.decode('latin-1')
       
        # CSV files - custom processing
        elif file_extension == '.csv':
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
                temp_file.write(file_obj.read())
                temp_path = temp_file.name
           
            try:
                return extract_text_from_csv(temp_path)
            finally:
                try:
                    os.unlink(temp_path)
                except:
                    pass
       
        # Other file types - save temporarily and process
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
            temp_file.write(file_obj.read())
            temp_path = temp_file.name
       
        try:
            if file_extension == '.pdf':
                return extract_text_from_pdf(temp_path)
            elif file_extension in ['.doc', '.docx']:
                return extract_text_from_word(temp_path)
            else:
                return f"Unsupported file type: {file_extension}"
        finally:
            # Clean up the temporary file
            try:
                os.unlink(temp_path)
            except:
                pass
   
    except Exception as e:
        logger.error(f"Error extracting text from document: {str(e)}")
        return f"Error extracting document text: {str(e)}"
 
def extract_text_from_csv(file_path):
    """Extract text from CSV files with better formatting."""
    try:
        import csv
       
        text = ""
        with open(file_path, 'r', newline='', encoding='utf-8') as file:
            try:
                csv_reader = csv.reader(file)
                headers = next(csv_reader)
               
                # Add headers
                text += "HEADERS: " + ", ".join(headers) + "\n\n"
               
                # Add rows with header labels (limit to first 100 rows for large files)
                for i, row in enumerate(csv_reader, 1):
                    if i > 100:
                        text += f"\n... [truncated, showing first 100 of more rows] ...\n"
                        break
                       
                    text += f"ROW {i}:\n"
                    for j, (header, value) in enumerate(zip(headers, row)):
                        text += f"  {header}: {value}\n"
                    text += "\n"
               
                return text
            except Exception as e:
                # Fallback to simple reading if CSV parsing fails
                file.seek(0)
                return file.read()
    except Exception as e:
        logger.error(f"Error extracting text from CSV: {str(e)}")
        return f"Error extracting CSV text: {str(e)}"
 
def extract_text_from_pdf(file_path):
    """Extract text from a PDF file."""
    try:
        try:
            import PyPDF2
        except ImportError:
            return "PyPDF2 library is not installed. Install with: pip install PyPDF2"
       
        text = ""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            num_pages = len(reader.pages)
           
            # Add document metadata
            text += f"PDF DOCUMENT: {num_pages} pages\n\n"
       
            for page_num in range(num_pages):
                page = reader.pages[page_num]
                page_text = page.extract_text()
               
                text += f"--- PAGE {page_num + 1} ---\n"
                text += page_text + "\n\n"
       
        return text
   
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        return f"Error extracting PDF text: {str(e)}"
 
def extract_text_from_word(file_path):
    """Extract text from a Word document."""
    try:
        # Try to import docx
        try:
            import docx
        except ImportError:
            return "python-docx library is not installed. Install with: pip install python-docx"
       
        # Load the document
        doc = docx.Document(file_path)
       
        # Extract metadata
        text = "WORD DOCUMENT\n\n"
       
        # Extract text from paragraphs with section headers
        current_section = "Main Document"
        text += f"SECTION: {current_section}\n"
       
        for para in doc.paragraphs:
            # Check if this might be a heading
            if para.style.name.startswith('Heading'):
                current_section = para.text
                text += f"\nSECTION: {current_section}\n"
           
            if para.text.strip():  # Only include non-empty paragraphs
                text += para.text + "\n"
       
        # Extract text from tables
        if doc.tables:
            text += "\nTABLES:\n"
            for i, table in enumerate(doc.tables):
                text += f"\nTable {i+1}:\n"
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text += row_text + "\n"
                text += "\n"
       
        return text
   
    except Exception as e:
        logger.error(f"Error extracting text from Word document: {str(e)}")
        return f"Error extracting Word document text: {str(e)}"